from requests import *
import pyttsx3
from customtkinter import *
import base64
from tkinter import *
from tkinter import messagebox
import os
from b64_imgs import *
from time import sleep
import docx # type: ignore
import pytesseract # type: ignore
from gpt4all import GPT4All #type: ignore


pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'





class Bw_funcoes:
    def obter_pesquisa(self):
        try:
            self.campo_busca = self.consulta_entry.get()
            match self.campo_busca:
                case '':
                    raise ValueError('Você precisa digitar algo!!!')
        except ValueError as erro:
            messagebox.showerror('Erro!', f'{erro}')

        else:
            try:
                # -=-=-=-=-=-=-=-=-=- Consumindo a api =-=-=-=-=-=-=-=-=-=-=-=-=-=

                engine = pyttsx3.init()
                model = GPT4All("Meta-Llama-3-8B-Instruct.Q4_0.gguf")
                with model.chat_session():
                    self.resposta = model.generate(self.campo_busca)
                # engine.say(self.resposta)
                # engine.runAndWait()

            except:
                messagebox.showerror('erro', 'GDP Trial expirou')
            else:
                # resposta=resposta.split()
                self.resposta_box.delete('0.0', END)
                self.resposta_box.insert('1.0', text=self.resposta)

    def cria_diretorio(self):
        self.dir_name = filedialog.asksaveasfilename(title='Escolha o local onde deseja criar a sua pasta '
                                                          'e dê um nome a ela!')
        try:
            pasta_existe = os.path.isdir(self.dir_name)
            if pasta_existe == False and not self.dir_name:
                raise ValueError('Você não escolheu um diretório, tente novamente!')
            elif pasta_existe:
                raise ValueError('Este diretório já existe, tente novamente!')
        except ValueError as erro:
                messagebox.showerror("ERRO!", f'{erro}')
        else:
            os.mkdir(self.dir_name)
            messagebox.showinfo("SUCESSO!", 'Sua pasta foi criada com sucesso!!')


    def gera_audio(self):
        self.nome_do_arquivo_audio = filedialog.asksaveasfilename(title='Digite um nome para o seu arquivo de audio e '
                                                             'escolha um local para salvá-lo!')
        if not self.nome_do_arquivo_audio:
            pass
        else:
            self.resposta = self.resposta_box.get('1.0', 'end')

            match len(self.resposta):
                case 1:
                    messagebox.showerror('erro', 'Você não realizou nenhuma pesquisa.')
                case _:
                    try:
                        self.resposta
                    except AttributeError:
                        messagebox.showerror('ERRO!',"Você não realizou nenhuma pesquisa!! Tente novamente.")
                    else:
                        engine = pyttsx3.init()
                        engine.save_to_file(f'{self.resposta}', f'{self.nome_do_arquivo_audio}.mp3')
                        engine.runAndWait()
                        messagebox.showinfo("SUCESSO!", f"O seu áudio foi gerado com sucesso na pasta escolhida! ")

    def gera_docx(self):

        # =-=-=-=-=- Pega o texto obtido pela pesquisa na janela de texto (1.0 indica o início do texto, end o fim!)
        self.text= self.resposta_box.get('1.0', 'end')

        match len(self.text):
            case 1:
                messagebox.showerror("ERRO", 'Você não realizou nenhuma pesquisa!')
            case _:
                # -=-==-=-=- Cria um arquivo txt com o texto obtido na janela de texto
                with open(file='./txt_file.txt', mode='w', encoding='utf-8') as tfile:
                    tfile.writelines(self.text)

                # -=-==-=-=- Cria um novo arquivo do Word
                self.arquivo = docx.Document()

                # -=-==-=-=- Lê o arquivo de texto e adiciona o conteúdo como um parágrafo no Word
                with open('./txt_file.txt', 'r', encoding='utf-8') as arquivo:
                    conteudo = arquivo.read()
                    self.arquivo.add_paragraph(conteudo)

                # -=-==-=-=- Salva o arquivo do Word
                self.novo_docx = filedialog.asksaveasfilename(title='Digite um nome para o seu arquivo word e '
                              'escolha um local para salvá-lo!')
                match self.novo_docx:
                    case '':
                        pass
                    case _:
                        self.arquivo.save(f'{self.novo_docx}.docx')
                        messagebox.showinfo('Sucesso', 'Arquivo Docx criado com sucesso!')


    def img_to_text(self):
        #=-=-=-=-=-=- Escolhe o arquivo em imagem para transformar em texto
        image_path = filedialog.askopenfilename(title='Escolha o arquivo de imagem que contém o texto desejado!',
                                                      filetypes=[("Formatos suportados: jpg, png, gif",
                                                                  "*.jpg;*.png;*.gif")])
        if image_path:
            # =-=-=-=- Converte a foto em texto
            texto = pytesseract.image_to_string(image_path, lang='por+eng')
            self.salvar_texto = filedialog.asksaveasfilename(title='Digite um nome para o seu arquivo de texto e '
                      'escolha um local para salvá-lo!')


            self.arquivo_doc = docx.Document() # <= Inicia o processo de criação do arquivo word
            self.arquivo_doc.add_paragraph(texto)
            self.arquivo_doc.save(f'{self.salvar_texto}.docx')


            messagebox.showinfo("Sucessso", 'Seu texto foi gerado com SUCESSO!')
        else:
            pass








