from mundo2.funcoes import *
from tqdm import tqdm
import os
import json
import docx
# -=-=-=-=-= "Paleta de cores" =-=-=-=-=-=-=-=-
os.system("")
cores = {
    'preto': '\033[1;30m',
    'vermelho': '\033[1;31m',
    'verde': '\033[1;32m',
    'vermelho claro': '\033[1;91m',
    'verde claro': '\033[1;92m',
    'negrito': '\033[;1m',
    'azul': '\033[1;34m',
    'amarelo': '\033[1;33m',
    'semcor': '\033[m'
}
# -=-=-=-=-=-=-=-=-=-= <INÍCIO DO PROGRAMA> =-=-=-=-=-=-=-=-=-=-=-=
print(f'{cores["azul"]}I-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^=I{cores["semcor"]}')
print(f'{cores["azul"]}I                          Bem vindo ao G.D.P  1.0 >>[GESTÃO DE PORTARIA]<<                                I{cores["semcor"]}')
print(f'{cores["azul"]}I     Feito por: Felipe Rodrigues Fonseca / contato: comunidadehawks@gmail.com / (55) 99724-6397           I{cores["semcor"]}')
print(f'{cores["azul"]}I-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^=I{cores["semcor"]}')

input(f'{cores["negrito"]}Aperte enter para continuar!{cores["semcor"]}')



#Criando banco de dados para chaves
statusChave = CriaBD('chaves.txt','nums.txt',)

#Criando banco de dados para ramais
statusRamal = CriaBD('NomeRamal.txt','NumRamal.txt',)

#Criando banco de dados para lista de afazeres
statusAfazeres = SimpleBD('Afazeres.txt',)

#Criando banco de dados para alterações no setor
statusAlteracoes = BdJson('Alt.json',)

#Criando dados para projetores
statusProjetores = BdJson('Projetores.json')

statusList = list()
statusList.append(statusChave)
statusList.append(statusRamal)
statusList.append(statusAfazeres)
statusList.append(statusAlteracoes)
statusList.append(statusProjetores)







temp1 = [] #lista para as chaves
temp2 = [] #lista para as chaves
temp3 = [] #lista para alterações no setor
temp4 = [] #lista para projetores


while True:
    if False in statusList:
        print(statusList)
        posicao = statusList.index(False)
        print(f'{cores["vermelho"]}HOUVE UM ERRO COM O BANCO DE DADOS!{cores["semcor"]}')
        match posicao+1:
            case 1:
                print('ERRO: BANCO PARA CHAVES!')
            case 2:
                print('ERRO: BANCO PARA RAMAL')
            case 3:
                print('ERRO: BANCO PARA AFAZERES!')
            case 4:
                print('ERRO: BANCO PARA ALTERAÇÕES')
            case 5:
                print('ERRO: BANCO PARA PROJETORES')
        input('Aperte enter para sair...')
        break
    else:
        print(f'{cores["negrito"]}BANCO DE DADOS STATUS: {cores["verde"]} [OK] {cores["semcor"]}')
    # -=-=-=-=-=-=-=-=-=-=-=-=-=  MENU PRINCIPAL-=-=-=-=-=-=-=-=-=-=-=-=-=-=-=-=
    opc_inicial = ''
    try:
        opc_inicial = menu_princ([f'ACESSAR MENU', 'UTILIZAR ASSISTENTE VIRTUAL (VÍRTUA)', 'SAIR'], 'MENU PRINCIPAL')
        print()
        if opc_inicial < 1 or opc_inicial > 3:
            print(f'{cores["negrito"]}{cores["vermelho claro"]}ERRO!{cores["semcor"]} OPÇÃO => {cores["vermelho claro"]}{opc_inicial}{cores["semcor"]} <= INVÁLIDA!!')
            continue
    except (TypeError, ValueError):
        mensagemerro()
    else:
        if opc_inicial == 1:
            # -=-=-=-=-=-=-=-=-=-=-=-=-=  MENU INICIAL -=-=-=-=-=-=-=-=-=-=-=-=-=-=-=-=
            while True:
                try:
                    menuin = menu(['CONTROLE DE CHAVES','CONTROLE DE RAMAL','LISTA DE AFAZERES','ALTERAÇÕES NO SETOR','CONTROLE DE PROJETORES','COBRANÇA DE MATERIAIS','VOLTAR AO MENU PRINCIPAL'],'MENU INICIAL')
                    if menuin < 1 or menuin > 7:
                        print(f'{cores["negrito"]}{cores["vermelho claro"]}ERRO!{cores["semcor"]} OPÇÃO => {cores["vermelho claro"]}{menuin}{cores["semcor"]} <= INVÁLIDA!!')
                        continue
                except (TypeError, ValueError):
                    mensagemerro()
                else:
                    if menuin == 7:
                        print(f'{cores["negrito"]}VOLTANDO AO MENU PRINCIPAL')
                        sleep(1.2)
                        print()
                        break
                    if menuin == 1:
                        # -=-=-=-=-=-=-=-=-=-=-=-=-=-=-=-=-=-=-= CÓDIGOS PARA CONTROLE DE CHAVES -=-=-=-=-=-=-=-=-=-=-=-=-=-=-=-=-=-=-=
                        while True:
                            try:
                                opcCH = menu([f'ADICIONAR CHAVE', 'REMOVER CHAVE','EDITAR CHAVE','CONSULTAR CADASTRO','LIMPAR CADASTRO','SAIR\n'],'MENU PARA CHAVES') # <========= CRIANDO O MENU
                                if opcCH < 1 or opcCH > 6:
                                    print(f'{cores["negrito"]}{cores["vermelho claro"]}ERRO!{cores["semcor"]} OPÇÃO => {cores["vermelho claro"]}{opcCH}{cores["semcor"]} <= INVÁLIDA!!')
                                    continue
                            except (TypeError, ValueError):
                                mensagemerro()
                            else:
                                if opcCH == 6:
                                    print(f'{cores["negrito"]}VOLTANDO AO MENU INICIAL...\033[m')
                                    sleep(1.2)
                                    print()
                                    break

                                    #Códigos para ADICIONAR chaves
                                elif opcCH == 1:
                                    print(f'Você escolheu a opção {cores["verde"]}{opcCH} (ADICIONAR CHAVE){cores["semcor"]}. Para continuar digite as seguintes informações:\n')
                                    while True:
                                        # Limpando as listas para que estejam aptas a receber os dados novamente impossibilitando duplicações de dados
                                        temp1.clear()
                                        temp2.clear()
                                        # abrindo arquivos em modo de leitura
                                        with open('chaves.txt', 'r') as chaves:
                                            dados = chaves.readlines()
                                            for nomes in dados:
                                                temp1.append(nomes.strip())

                                        with open('nums.txt', 'r') as nums:
                                            dados = nums.readlines()
                                            for num in dados:
                                                temp2.append(num.strip())

                                        # inicio da estrutura de cadastro
                                        try:
                                            chave = input('Digite o nome da chave que será cadastrada ou digite "sair" para voltar ao menu: ').strip()
                                            if not chave:
                                                raise ValueError('ERRO VOCÊ PRECISA DIGITAR ALGO!')
                                            elif chave == 'sair' or chave == 'Sair':
                                                print('VOLTANDO...')
                                                chave = ''
                                                sleep(1.0)
                                                break
                                            elif chave in temp1:
                                                raise ValueError('ERRO CHAVE JÁ CADASTRADA!')
                                        except ValueError as Erro:
                                            print(Erro)
                                            continue
                                        else:
                                            temp1.append(chave)
                                        while ValueError:
                                            cont = False
                                            try:
                                                num = input('Digite o número da chave: ').strip()
                                                if not num:
                                                    raise ValueError('ERRO, VOCÊ PRECISA DIGITAR ALGO!')
                                                elif num in temp2:
                                                    raise ValueError('ERRO, NÚMERO USADO!!!')
                                            except ValueError as Erro:
                                                print(f'{cores["vermelho"]}{Erro}{cores["semcor"]}')
                                            else:
                                                temp2.append(num)
                                                print()
                                                # abrindo arquivos em modo de escrita
                                                with open('chaves.txt', 'w') as bdchaves:
                                                    for dados in temp1:
                                                        bdchaves.write(f'{dados}\n')
                                                with open('nums.txt', 'w') as bdnums:
                                                    for dados2 in temp2:
                                                        bdnums.write(f'{dados2}\n')
                                                print(f'{cores["verde"]}Adicionado a chave:{cores["semcor"]} {cores["negrito"]}{chave}{cores["semcor"]} {cores["verde"]}no número:{cores["semcor"]} {cores["negrito"]}{num}!{cores["semcor"]}')
                                                cont = True
                                                if cont:
                                                    break

                                #  Códigos para REMOVER chaves

                                elif opcCH == 2:
                                    print(f'Você escolheu a opção {cores["verde"]}{opcCH} (REMOVER CHAVE){cores["semcor"]}. Para continuar digite as seguintes informações:\n')
                                    while True:
                                        temp1.clear()
                                        temp2.clear()
                                        with open('chaves.txt', 'r') as DadosChaves:
                                            dados = DadosChaves.readlines()
                                            for chave in dados:
                                                temp1.append(chave.strip())

                                        with open('nums.txt', 'r') as DadosNums:
                                            dnums = DadosNums.readlines()
                                            for n in dnums:
                                                temp2.append(n.strip())
                                        try:
                                            chRem = input('Digite qual chave deseja remover ou "sair" para voltar ao menu: ')
                                            if chRem == 'sair' or chRem == 'Sair':
                                                print('Voltando ao Menu...')
                                                break
                                            elif not chRem:
                                                raise ValueError('ERRO! VOCÊ PRECISA DIGITAR ALGO!')
                                            elif chRem not in temp1:
                                                raise ValueError('ERRO, chave não cadastrada!')
                                            elif not chRem:
                                                raise ValueError('ERRO! VOCÊ PRECISA DIGITAR ALGO!')
                                        except ValueError as erro:
                                            print(f'{cores["vermelho"]}{erro}{cores["semcor"]}')
                                            continue
                                        else:
                                            i = temp1.index(chRem)
                                            temp1.remove(chRem)
                                            print(f'REMOVIDO COM SUCESSO A CHAVE {chRem} COM O Nº:{temp2[i]}. ')
                                            temp2.pop(i)
                                            with open('chaves.txt', 'w') as dChaves:
                                                for chave in temp1:
                                                    dChaves.write(f'{chave}\n')

                                            with open('nums.txt', 'w') as dNumeros:
                                                for nume in temp2:
                                                    dNumeros.write(f'{nume}\n')

                                # Códigos para EDITAR chaves
                                elif opcCH == 3:
                                    print(f'Você escolheu a opção {cores["verde"]}{opcCH} (EDITAR CHAVE){cores["semcor"]}. Para continuar escolha uma opção:\n')
                                    while True:
                                        try:
                                            opcCHEDIT = menu(['EDITAR O NOME DA CHAVE', 'EDITAR NÚMERO DA CHAVE','VOLTAR AO MENU ANTERIOR'],'OPÇÕES PARA EDITAR CHAVES')
                                            if opcCHEDIT < 1 or opcCHEDIT > 3:
                                                print(f'{cores["negrito"]}{cores["vermelho claro"]}ERRO!{cores["semcor"]} OPÇÃO => {cores["vermelho claro"]}{opcCHEDIT}{cores["semcor"]} <= INVÁLIDA!!')
                                                continue
                                        except (ValueError, TypeError):
                                            mensagemerro()
                                        else:
                                            if opcCHEDIT == 3:
                                                print(f'{cores["negrito"]}VOLTANDO AO MENU ANTERIOR...\033[m')
                                                sleep(1.2)
                                                print()
                                                break
                                            #CÓDIGOS PARA EDITAR NOME DA CHAVE
                                            elif opcCHEDIT == 1:
                                                print(f'Você escolheu a opção {cores["verde"]}{opcCHEDIT} (EDITAR O NOME DA CHAVE){cores["semcor"]}. Para continuar digite as informações abaixo:\n')
                                                while True:
                                                    temp1.clear()
                                                    with open('chaves.txt', 'r') as DadosChave:
                                                        dados = DadosChave.readlines()
                                                        for chave in dados:
                                                            temp1.append(chave.strip())
                                                    try:
                                                        chEdit = input('Digite qual chave será editada ou sair para voltar ao menu: ')
                                                        if chEdit == 'sair' or chEdit == 'Sair':
                                                            print('Voltando ao Menu...')
                                                            break
                                                        elif not chEdit:
                                                            raise ValueError('ERRO, VOCÊ PRECISA DIGITAR ALGO!')
                                                        elif chEdit not in temp1:
                                                            raise ValueError('ERRO, CHAVE NÃO CADASTRADA!')
                                                    except ValueError as erro:
                                                        print(f'{cores["vermelho"]}{erro}{cores["semcor"]}')
                                                        continue
                                                    else:
                                                        i = temp1.index(chEdit)
                                                    cont = False
                                                    while True:
                                                        try:
                                                            novoNome = input('Digite o novo nome para a chave ou sair para voltar ao menu: ')
                                                            if novoNome == 'sair' or novoNome == 'Sair':
                                                                novoNome = ''
                                                                print('SAINDO...')
                                                                break
                                                            elif novoNome in temp1:
                                                                raise ValueError(f'ERRO, o nome {novoNome} já está sendo utilizado!')
                                                            elif not novoNome:
                                                                raise ValueError('ERRO, VOCÊ PRECISA DIGITAR ALGO!')
                                                        except ValueError as erro:
                                                            print(f'{cores["vermelho"]}{erro}{cores["semcor"]}')
                                                            continue
                                                        else:
                                                            temp1[i] = novoNome
                                                            with open('chaves.txt', 'w') as DadosChave:
                                                                for chave in temp1:
                                                                    DadosChave.write(f'{chave}\n')
                                                            print(f'Chave {chEdit} alterada com SUCESSO para -> {novoNome} <-')
                                                            cont = True
                                                            if cont:
                                                                break

                                            #CÓDIGO PARA EDITAR O NÚMERO DA CHAVE
                                            elif opcCHEDIT == 2:
                                                print(f'Você escolheu a opção {cores["verde"]}{opcCHEDIT} (EDITAR O NÚMERO DA CHAVE){cores["semcor"]}. Para continuar digite as informações abaixo:\n')
                                                while True:
                                                    temp1.clear()
                                                    temp2.clear()
                                                    with open('chaves.txt', 'r') as bdChaves:
                                                        dados = bdChaves.readlines()
                                                        for chave in dados:
                                                            temp1.append(chave.strip())
                                                    with open('nums.txt', 'r') as bdNums:
                                                        dados = bdNums.readlines()
                                                        for num in dados:
                                                            temp2.append(num.strip())
                                                    try:
                                                        numEdit = input('Digite qual chave deseja editar o número ou digite "sair" para voltar ao menu: ')
                                                        if numEdit == 'sair' or numEdit == 'Sair':
                                                            print('Voltando...')
                                                            break
                                                        elif not numEdit:
                                                            raise ValueError('ERRO, VOCÊ PRECISA DIGITAR ALGO!')
                                                        elif numEdit not in temp1:
                                                            raise ValueError('ERRO, chave não encontrada!')
                                                    except ValueError as erro:
                                                        print(f'{cores["vermelho"]}{erro}{cores["semcor"]}')
                                                        continue
                                                    else:
                                                        cont = False
                                                        while True:
                                                            try:
                                                                novoNum = input('Digite o novo número: ')
                                                                if novoNum in temp2:
                                                                    raise ValueError(f'ERRO, o número {novoNum} já está sendo utilizado!')
                                                                elif not novoNum:
                                                                    raise ValueError('ERRO, VOCÊ PRECISA DIGITAR ALGO!')
                                                            except ValueError as Erro:
                                                                print(f'{cores["vermelho"]}{Erro}{cores["semcor"]}')
                                                                continue
                                                            else:
                                                                pos = temp1.index(numEdit)
                                                                print(f'O número da chave {numEdit} que era {temp2[pos]} foi alterado para {novoNum} COM SUCESSO!')
                                                                temp2[pos] = novoNum
                                                                with open('nums.txt', 'w') as bdNums:
                                                                    for num in temp2:
                                                                        bdNums.write(f'{num}\n')
                                                                cont = True
                                                                if cont:
                                                                    break
                                elif opcCH == 4:
                                    print(f'Você escolheu a opção {cores["verde"]}{opcCH} (CONSULTAR CADASTRO){cores["semcor"]}.\n')
                                    print(f'{cores["azul"]}I-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^=I{cores["semcor"]}')
                                    print(f'{cores["azul"]}I                          >>[CHAVES CADASTRADAS ATÉ O MOMENTO]<<                                          I{cores["semcor"]}')
                                    print(f'{cores["azul"]}I                                                                                                          I{cores["semcor"]}')
                                    print(f'{cores["azul"]}I-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^=I{cores["semcor"]}')
                                    with open('chaves.txt','r') as BancoChaves:
                                        dados = BancoChaves.readlines()
                                        with open('nums.txt','r') as BancoNums:
                                            dadosnum = BancoNums.readlines()
                                            for chave in dados:
                                                pos = dados.index(chave)
                                                print(f'{chave.strip()} - Nº: {dadosnum[pos].strip()}')

                                    input('PRESSIONE ENTER PARA VOLTAR!')
                                    print('VOLTANDO AO MENU...')
                                    sleep(1.0)

                                elif opcCH == 5:
                                    print(f'Você escolheu a opção {cores["verde"]}{opcCH} (LIMPAR CADASTRO){cores["semcor"]}\n')
                                    pinta('vermelho','ATENÇÃO! ESTA OPÇÃO IRÁ APAGAR TODOS OS DADOS SALVOS DE UMA VEZ SÓ!')
                                    pinta('vermelho','APÓS APAGAR, OS DADOS NÃO PODERÃO SER RESGATADOS NOVAMENTE!')
                                    while True:
                                        try:
                                            opcoes = ['n','s']
                                            prosseguir = input('Deseja continuar? [s]- para continuar / [n] para cancelar e voltar ao menu: ').lower()
                                            if not prosseguir:
                                                raise ValueError('ERRO VOCÊ PRECISA ESCOLHER UMA DAS OPÇÕES!')
                                            elif prosseguir not in opcoes:
                                                raise ValueError(f'ERRO,OPÇÃO -> {prosseguir} <- INVÁLIDA!')
                                        except ValueError as erro:
                                            print(f'{cores["vermelho"]}{erro}{cores["semcor"]}')
                                            continue
                                        else:
                                            if prosseguir in opcoes and prosseguir == 'n':
                                                pinta('negrito','VOLTANDO AO MENU DE CHAVES')
                                                sleep(2.0)
                                                break
                                            elif prosseguir in opcoes and prosseguir == 's':
                                                print(f'{cores["negrito"]}ESVAZIANDO CADASTRO DE CHAVES...{cores["semcor"]}')
                                                with open('chaves.txt','r') as bdchaves:
                                                    dados = bdchaves.readlines()
                                                    if len(dados) == 0:
                                                        print(f'{cores["verde"]}O CADASTRO JÁ ESTÁ VAZIO!{cores["semcor"]}')
                                                        sleep(2.0)
                                                        break
                                                    else:
                                                        for dado in tqdm(dados):
                                                            sleep(0.5)
                                                with open('chaves.txt', 'w') as bdchaves:
                                                    print()
                                                with open('nums.txt', 'w') as bdnums:
                                                    print(f'{cores["verde"]}CADASTRO ESVAZIADO COM SUCESSO!!!{cores["semcor"]}')
                                                sleep(2.0)
                                                break

                    #INICIO DA ESTRUTURA DE CONTROLE DE RAMAIS
                    elif menuin == 2:
                        while True:
                            try:
                                opcRA = menu([f'ADICIONAR RAMAL', 'REMOVER RAMAL', 'EDITAR RAMAL', 'CONSULTAR CADASTRO','LIMPAR CADASTRO','SAIR\n'], 'MENU PARA RAMAIS')  # <========= CRIANDO O MENU
                                if opcRA < 1 or opcRA > 6:
                                    print(f'{cores["negrito"]}{cores["vermelho claro"]}ERRO!{cores["semcor"]} OPÇÃO => {cores["vermelho claro"]}{opcRA}{cores["semcor"]} <= INVÁLIDA!!')
                                    continue
                            except (TypeError, ValueError):
                                mensagemerro()
                            else:
                                if opcRA == 6:
                                    print(f'{cores["negrito"]}VOLTANDO AO MENU INICIAL...\033[m')
                                    sleep(1.2)
                                    print()
                                    break

                                # -=-=-=-=-=-=-=-=-=-=-=-=-=-=-=-=-=-=-= CÓDIGOS PARA CONTROLE DE RAMAIS -=-=-=-=-=-=-=-=-=-=-=-=-=-=-=-=-=-=-=

                                # Códigos para ADICIONAR RAMAIS
                                elif opcRA == 1:
                                    print( f'Você escolheu a opção {cores["verde"]}{opcRA} (ADICIONAR RAMAL){cores["semcor"]}. Para continuar digite as seguintes informações:\n')
                                    while True:
                                        # Limpando as listas para que estejam aptas a receber os dados novamente impossibilitando duplicações de dados
                                        temp1.clear()
                                        temp2.clear()
                                        # abrindo arquivos em modo de leitura
                                        with open('NomeRamal.txt', 'r') as rama:
                                            dados = rama.readlines()
                                            for ramal in dados:
                                                temp1.append(ramal.strip())

                                        with open('NumRamal.txt', 'r') as nums:
                                            dados = nums.readlines()
                                            for num in dados:
                                                temp2.append(num.strip())

                                        # inicio da estrutura de cadastro
                                        try:
                                            NovoRam = input('Digite a descrição do Ramal que será cadastrado ou digite "sair" para voltar ao menu: ').strip()
                                            if not NovoRam:
                                                raise ValueError('ERRO VOCÊ PRECISA DIGITAR ALGO!')
                                            elif NovoRam == 'sair' or NovoRam == 'Sair':
                                                print('VOLTANDO...')
                                                NovoRam = ''
                                                sleep(1.0)
                                                break
                                            elif NovoRam in temp1:
                                                raise ValueError('ERRO RAMAL JÁ CADASTRADO!')
                                        except ValueError as Erro:
                                            print(f'{cores["vermelho"]}{Erro}{cores["semcor"]}')
                                            continue
                                        else:
                                            temp1.append(NovoRam)
                                        while ValueError:
                                            cont = False
                                            try:
                                                num = input('Digite o número do Ramal: ').strip()
                                                if not num:
                                                    raise ValueError('ERRO, VOCÊ PRECISA DIGITAR ALGO!')
                                                elif num in temp2:
                                                    raise ValueError('ERRO, NÚMERO USADO!!!')
                                            except ValueError as Erro:
                                                print(f'{cores["vermelho"]}{Erro}{cores["semcor"]}')
                                            else:
                                                temp2.append(num)
                                                print()
                                                # abrindo arquivos em modo de escrita
                                                with open('NomeRamal.txt', 'w') as bdram:
                                                    for dados in temp1:
                                                        bdram.write(f'{dados}\n')
                                                with open('NumRamal.txt', 'w') as bdnums:
                                                    for dados2 in temp2:
                                                        bdnums.write(f'{dados2}\n')
                                                print(f'Adicionado o Ramal: {NovoRam} correspondente ao número: {num}!')
                                                cont = True
                                                if cont:
                                                    break

                                #  CÓDIGOS PARA REMOVER RAMAIS
                                elif opcRA == 2:
                                    print(f'Você escolheu a opção {cores["verde"]}{opcRA} (REMOVER RAMAL){cores["semcor"]}. Para continuar digite as seguintes informações:\n')
                                    while True:
                                        temp1.clear()
                                        temp2.clear()
                                        with open('NomeRamal.txt', 'r') as DadosRam:
                                            dados = DadosRam.readlines()
                                            for ramal in dados:
                                                temp1.append(ramal.strip())

                                        with open('NumRamal.txt', 'r') as DadosNums:
                                            dnums = DadosNums.readlines()
                                            for n in dnums:
                                                temp2.append(n.strip())
                                        try:
                                            RamalRem = input('Digite qual Ramal deseja remover ou "sair" para voltar ao menu: ')
                                            if RamalRem == 'sair' or RamalRem == 'Sair':
                                                print('Voltando ao Menu...')
                                                break
                                            elif not RamalRem:
                                                raise ValueError('ERRO! VOCÊ PRECISA DIGITAR ALGO!')
                                            elif RamalRem not in temp1:
                                                raise ValueError('ERRO, RAMAL NÃO CADASTRADO!')
                                        except ValueError as erro:
                                            print(f'{cores["vermelho"]}{erro}{cores["semcor"]}')
                                            continue
                                        else:
                                            i = temp1.index(RamalRem)
                                            temp1.remove(RamalRem)
                                            print(f'REMOVIDO COM SUCESSO O RAMAL {RamalRem} COM O Nº:{temp2[i]}. ')
                                            temp2.pop(i)
                                            with open('NomeRamal.txt', 'w') as dRamal:
                                                for ramal in temp1:
                                                    dRamal.write(f'{ramal}\n')

                                            with open('NumRamal.txt', 'w') as dNumeros:
                                                for nume in temp2:
                                                    dNumeros.write(f'{nume}\n')

                                # CÓDIGOS PARA EDITAR RAMAIS
                                elif opcRA == 3:
                                    print(f'Você escolheu a opção {cores["verde"]}{opcRA} (EDITAR RAMAL){cores["semcor"]}. Para continuar escolha uma opção:\n')
                                    while True:
                                        try:
                                            opcRAEDIT = menu(['EDITAR O NOME DO RAMAL', 'EDITAR O NÚMERO DO RAMAL', 'VOLTAR AO MENU ANTERIOR'],'OPÇÕES PARA EDITAR RAMAL')
                                            if opcRAEDIT < 1 or opcRAEDIT > 3:
                                                print(
                                                    f'{cores["negrito"]}{cores["vermelho claro"]}ERRO!{cores["semcor"]} OPÇÃO => {cores["vermelho claro"]}{opcRAEDIT}{cores["semcor"]} <= INVÁLIDA!!')
                                                continue
                                        except (TypeError, ValueError):
                                            mensagemerro()
                                        else:
                                            if opcRAEDIT == 3:
                                                print(f'{cores["negrito"]}VOLTANDO AO MENU ANTERIOR...\033[m')
                                                sleep(1.2)
                                                print()
                                                break

                                            # CÓDIGOS PARA EDITAR NOME DO RAMAL
                                            elif opcRAEDIT == 1:
                                                print(f'Você escolheu a opção {cores["verde"]}{opcRAEDIT} (EDITAR O NOME DO RAMAL){cores["semcor"]}. Para continuar digite as informações abaixo:\n')
                                                while True:
                                                    temp1.clear()
                                                    with open('NomeRamal.txt', 'r') as DadosRam:
                                                        dados = DadosRam.readlines()
                                                        for ramal in dados:
                                                            temp1.append(ramal.strip())
                                                    try:
                                                        raEdit = input('Digite qual Ramal será editado ou sair para voltar ao menu: ')
                                                        if raEdit == 'sair' or raEdit == 'Sair':
                                                            print('Voltando ao Menu...')
                                                            break
                                                        elif not raEdit:
                                                            raise ValueError('ERRO, VOCÊ PRECISA DIGITAR ALGO!')
                                                        elif raEdit not in temp1:
                                                            raise ValueError('ERRO, RAMAL NÃO CADASTRADO!')
                                                    except ValueError as erro:
                                                        print(f'{cores["vermelho"]}{erro}{cores["semcor"]}')
                                                        continue
                                                    else:
                                                        i = temp1.index(raEdit)
                                                    cont = False
                                                    while True:
                                                        try:
                                                            novoNome = input('Digite o novo nome para o Ramal ou sair para voltar ao menu: ')
                                                            if novoNome == 'sair' or novoNome == 'Sair':
                                                                novoNome = ''
                                                                print('SAINDO...')
                                                                break
                                                            elif novoNome in temp1:
                                                                raise ValueError(
                                                                    f'ERRO, o nome {novoNome} já está sendo utilizado!')
                                                            elif not novoNome:
                                                                raise ValueError('ERRO, VOCÊ PRECISA DIGITAR ALGO!')
                                                        except ValueError as erro:
                                                            print(f'{cores["vermelho"]}{erro}{cores["semcor"]}')
                                                            continue
                                                        else:
                                                            temp1[i] = novoNome
                                                            with open('NomeRamal.txt', 'w') as DadosRam:
                                                                for ramal in temp1:
                                                                    DadosRam.write(f'{ramal}\n')
                                                            print(f'O Ramal {raEdit} foi alterado com SUCESSO para -> {novoNome} <-')
                                                            cont = True
                                                            if cont:
                                                                break

                                            # CÓDIGO PARA EDITAR O NÚMERO DO RAMAL
                                            elif opcRAEDIT == 2:
                                                print(f'Você escolheu a opção {cores["verde"]}{opcRAEDIT} (EDITAR O NÚMERO DO RAMAL){cores["semcor"]}. Para continuar digite as informações abaixo:\n')
                                                while True:
                                                    temp1.clear()
                                                    temp2.clear()
                                                    with open('NomeRamal.txt', 'r') as bdRams:
                                                        dados = bdRams.readlines()
                                                        for ramal in dados:
                                                            temp1.append(ramal.strip())
                                                    with open('NumRamal.txt', 'r') as bdNums:
                                                        dados = bdNums.readlines()
                                                        for num in dados:
                                                            temp2.append(num.strip())
                                                    try:
                                                        RamEdit = input('Digite qual Ramal deseja editar o número ou digite "sair" para voltar ao menu: ')
                                                        if RamEdit == 'sair' or RamEdit == 'Sair':
                                                            print(f'{cores["negrito"]}Voltando...{cores["semcor"]}')
                                                            break
                                                        elif not RamEdit:
                                                            raise ValueError('ERRO, VOCÊ PRECISA DIGITAR ALGO!')
                                                        elif RamEdit not in temp1:
                                                            raise ValueError('ERRO, RAMAL NÃO ENCONTRADO!')
                                                    except ValueError as erro:
                                                        print(f'{cores["vermelho"]}{erro}{cores["semcor"]}')
                                                        continue
                                                    else:
                                                        cont = False
                                                        while True:
                                                            try:
                                                                novoNum = input('Digite o novo número: ')
                                                                if novoNum in temp2:
                                                                    raise ValueError(f'ERRO, o número {novoNum} já está sendo utilizado!')
                                                                elif not novoNum:
                                                                    raise ValueError('ERRO, VOCÊ PRECISA DIGITAR ALGO!')
                                                            except ValueError as Erro:
                                                                print(f'{cores["vermelho"]}{Erro}{cores["semcor"]}')
                                                                continue
                                                            else:
                                                                pos = temp1.index(RamEdit)
                                                                print(f'O número do Ramal {RamEdit} que era {temp2[pos]} foi alterado para {novoNum} COM SUCESSO!')
                                                                temp2[pos] = novoNum
                                                                with open('NumRamal.txt', 'w') as bdNums:
                                                                    for num in temp2:
                                                                        bdNums.write(f'{num}\n')
                                                                cont = True
                                                                if cont:
                                                                    break
                                elif opcRA == 4:
                                    print(f'Você escolheu a opção {cores["verde"]}{opcRA} (CONSULTAR CADASTRO){cores["semcor"]}.\n')
                                    print(f'{cores["azul"]}I-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^=I{cores["semcor"]}')
                                    print(f'{cores["azul"]}I                          >>[RAMAIS CADASTRADOS ATÉ O MOMENTO]<<                                          I{cores["semcor"]}')
                                    print(f'{cores["azul"]}I                                                                                                          I{cores["semcor"]}')
                                    print(f'{cores["azul"]}I-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^=I{cores["semcor"]}')
                                    with open('NomeRamal.txt', 'r') as BancoRamais:
                                        dados = BancoRamais.readlines()
                                        with open('NumRamal.txt', 'r') as BancoNums:
                                            dadosnum = BancoNums.readlines()
                                            for ramal in dados:
                                                pos = dados.index(ramal)
                                                print(f'{ramal.strip()} - Nº: {dadosnum[pos].strip()}')

                                    input('PRESSIONE ENTER PARA VOLTAR!')
                                    print('VOLTANDO AO MENU...')
                                    sleep(1.0)
                                elif opcRA == 5:
                                    print(f'Você escolheu a opção {cores["verde"]}{opcRA} (LIMPAR CADASTRO){cores["semcor"]}\n')
                                    pinta('vermelho','ATENÇÃO! ESTA OPÇÃO IRÁ APAGAR TODOS OS DADOS SALVOS DE UMA VEZ SÓ!')
                                    pinta('vermelho', 'APÓS APAGAR, OS DADOS NÃO PODERÃO SER RESGATADOS NOVAMENTE!')
                                    while True:
                                        try:
                                            opcoes = ['n', 's']
                                            prosseguir = input('Deseja continuar? [s]- para continuar / [n] para cancelar e voltar ao menu: ').lower()
                                            if not prosseguir:
                                                raise ValueError('ERRO VOCÊ PRECISA ESCOLHER UMA DAS OPÇÕES!')
                                            elif prosseguir not in opcoes:
                                                raise ValueError(f'ERRO,OPÇÃO {cores["vermelho"]}-> {prosseguir} <-{cores["semcor"]} INVÁLIDA!')
                                        except ValueError as erro:
                                            print(f'{cores["vermelho"]}{erro}{cores["semcor"]}')
                                            continue
                                        else:
                                            if prosseguir in opcoes and prosseguir == 'n':
                                                pinta('negrito', 'VOLTANDO AO MENU DE RAMAIS')
                                                sleep(2.0)
                                                break
                                            elif prosseguir in opcoes and prosseguir == 's':
                                                print(f'{cores["negrito"]}ESVAZIANDO CADASTRO DE RAMAIS!{cores["semcor"]}')
                                                with open('NomeRamal.txt','r') as bdramal:
                                                    ramais = bdramal.readlines()
                                                    if len(ramais) == 0:
                                                        print(f'{cores["verde"]}O CADASTRO JÁ ESTÁ VAZIO!{cores["semcor"]}')
                                                        sleep(2.0)
                                                        break
                                                    else:
                                                        for chave in tqdm(ramais):
                                                            sleep(0.5)

                                                with open('NomeRamal.txt', 'w') as bdchaves:
                                                    print()
                                                with open('NumRamal.txt', 'w') as bdnums:
                                                    pinta('verde','CADASRO ESVAZIADO COM SUCESSO!')
                                                sleep(2.0)
                                                break
                        # CODIGOS PARA LISTA DE AFAZERES
                    elif menuin == 3:
                        while True:
                            try:
                                menuafa = menu(['ADICIONAR TAREFA', 'REMOVER TAREFA', 'CONSULTAR TAREFAS','LIMPAR LISTA DE AFAZERES','VOLTAR AO MENU ANTERIOR'],'OPÇÕES PARA LISTA DE AFAZERES')
                                if menuafa < 1 or menuafa > 5:
                                    print(f'{cores["negrito"]}{cores["vermelho claro"]}ERRO!{cores["semcor"]} OPÇÃO => {cores["vermelho claro"]}{menuafa}{cores["semcor"]} <= INVÁLIDA!!')
                                    continue
                            except (TypeError, ValueError):
                                mensagemerro()
                            else:
                                if menuafa == 5:
                                    pinta('negrito','VOLTANDO AO MENU INICIAL...')
                                    sleep(1.0)
                                    break
                                #INICIO DA ESTRUTURA DE ADICIONAR TAREFAS
                                elif menuafa == 1:
                                    print(f'Você escolheu a opção {cores["verde"]}{menuafa} (ADICIONAR TAREFA){cores["semcor"]}. Para continuar digite as informações abaixo:\n')
                                    while True:
                                        temp1.clear()
                                        with open('Afazeres.txt','r') as AfaDados:
                                            afazeres = AfaDados.readlines()
                                            for tarefa in afazeres:
                                                temp1.append(tarefa.strip())
                                        try:
                                            novaTarefa = input('Digite a Tarefa a ser lembrada juntamente com a data ou digite sair para voltar: ')
                                            if not novaTarefa:
                                                raise ValueError('ERRO, VOCÊ PRECISA DIGITAR ALGO!')
                                            elif novaTarefa in temp1:
                                                raise ValueError(f'ERRO, {novaTarefa} já está cadastrada!')
                                            elif novaTarefa == 'sair' or novaTarefa == 'Sair':
                                                pinta('negrito','VOLTANDO AO MENU ANTERIOR...')
                                                sleep(1.0)
                                                break
                                        except ValueError as erro:
                                            print(f'{cores["vermelho"]}{erro}{cores["semcor"]}')
                                            continue
                                        else:
                                            temp1.append(novaTarefa)
                                            with open('Afazeres.txt','w') as bdAfa:
                                                for tarefa in temp1:
                                                    bdAfa.write(f'{tarefa}\n')
                                            print(f'ATIVIDADE: {novaTarefa} cadastrada com SUCESSO!')

                                elif menuafa == 2:

                                    #INICIO DA ESTRUTURA DE REMOÇÃO DE TAREFAS
                                    print(f'Você escolheu a opção {cores["verde"]}{menuafa} (REMOVER TAREFA){cores["semcor"]}. Para continuar digite as informações abaixo:\n')
                                    while True:
                                        temp1.clear()
                                        with open('Afazeres.txt', 'r') as bdafa:
                                            dados = bdafa.readlines()
                                            for tarefas in dados:
                                                temp1.append(tarefas.strip())
                                        print(f'{cores["negrito"]} \nTAREFAS CADASTRADAS ATÉ O MOMENTO {cores["semcor"]}\n')
                                        for cod,tarefas in enumerate(temp1):
                                            print(f'CÓD: {cores["negrito"]}[{cod+1}]{cores["semcor"]} - Tarefa: {tarefas}')
                                        print()
                                        try:
                                            remover = input('Digite o Código da tarefa que deseja remover da lista ou "0" para voltar ao menu: ')
                                            if remover == '0':
                                                pinta('negrito','VOLTANDO AO MENU ANTERIOR...')
                                                sleep(1.5)
                                                break
                                            elif not remover:
                                                raise ValueError('ERRO, VOCÊ PRECISA DIGITAR ALGO!')
                                        except ValueError as Erro:
                                            print(f'{cores["vermelho"]}{Erro}{cores["semcor"]}')
                                            continue
                                        else:
                                            remover = int(remover)
                                            if remover < 0 or remover > len(temp1):
                                                print('ERRO, TAREFA NÃO CADASTRADA!')
                                                continue
                                            temp1.pop(remover-1)
                                            with open('Afazeres.txt', 'w') as bdafa:
                                                for tarefas in temp1:
                                                    bdafa.write(f'{tarefas}\n')
                                            print(f'A atividade {remover} foi excluída com SUCESSO!')
                                elif menuafa == 3:

                                    # INICIO DA ESTRUTURA DE CONSULTAR TAREFAS
                                    temp1.clear()
                                    print(f'Você escolheu a opção {cores["verde"]}{menuafa} (CONSULTAR TAREFAS){cores["semcor"]}\n')
                                    print(f'{cores["azul"]}I-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^=I{cores["semcor"]}')
                                    print(f'{cores["azul"]}I                               >>[TAREFAS CADASTRADAS ATÉ O MOMENTO]<<                                    I{cores["semcor"]}')
                                    print(f'{cores["azul"]}I                                                                                                          I{cores["semcor"]}')
                                    print(f'{cores["azul"]}I-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^-=^=I{cores["semcor"]}')
                                    with open('Afazeres.txt', 'r') as afadados:
                                        dados = afadados.readlines()
                                        for tarefa in dados:
                                            temp1.append(tarefa.strip())
                                    print()
                                    for indice,tarefa in enumerate(temp1):
                                        print(f'{indice+1} - {tarefa}')
                                    print()
                                    input('Aperte "ENTER" para voltar ao menu.')

                                elif menuafa == 4:
                                    print(f'Você escolheu a opção {cores["verde"]}{menuafa} (LIMPAR CADASTRO){cores["semcor"]}\n')
                                    pinta('vermelho','ATENÇÃO! ESTA OPÇÃO IRÁ APAGAR TODOS OS DADOS SALVOS DE UMA VEZ SÓ!')
                                    pinta('vermelho','APÓS APAGAR, OS DADOS NÃO PODERÃO SER RESGATADOS NOVAMENTE!')
                                    while True:
                                        try:
                                            opcoes = ['n', 's']
                                            prosseguir = input(
                                                'Deseja continuar? [s]- para continuar / [n] para cancelar e voltar ao menu: ').lower()
                                            if not prosseguir:
                                                raise ValueError('ERRO VOCÊ PRECISA ESCOLHER UMA DAS OPÇÕES!')
                                            elif prosseguir not in opcoes:
                                                raise ValueError(f'ERRO,OPÇÃO {cores["vermelho"]}-> {prosseguir} <-{cores["semcor"]} INVÁLIDA!')
                                        except ValueError as erro:
                                            print(f'{cores["vermelho"]}{erro}{cores["semcor"]}')
                                            continue
                                        else:
                                            if prosseguir in opcoes and prosseguir == 'n':
                                                pinta('negrito', 'VOLTANDO AO MENU LISTA DE AFAZERES')
                                                sleep(2.0)
                                                break
                                            elif prosseguir in opcoes and prosseguir == 's':
                                                print(f'{cores["negrito"]}ESVAZIANDO CADASTRO DE AFAZERES...{cores["semcor"]}')

                                                with open('Afazeres.txt', 'r') as bdafa:
                                                    tarefas = bdafa.readlines()
                                                    if len(tarefas) == 0:
                                                        print(f'{cores["verde"]}O CADASTRO JÁ ESTÁ VAZIO!{cores["semcor"]}')
                                                        sleep(2.0)
                                                        break
                                                    else:
                                                        for dado in tqdm(tarefas):
                                                            sleep(0.5)

                                                with open('Afazeres.txt', 'w') as bdafa:
                                                    print(f'{cores["verde"]}LISTA ESVAZIADA COM SUCESSO!!!{cores["semcor"]}')
                                                sleep(2.0)
                                                break
                    elif menuin == 4:
                        #Códigos para controle de alterações=================
                        while True:
                            try:
                                escolha = menu(['Adicionar Alteração','Consultar Alterações','Remover Alteração',
                                                'Limpar Dados','Enviar dados por E-mail','Voltar ao Menu Inicial'],'ALTERAÇÕES NO SETOR')
                                if escolha < 1 or escolha > 6:
                                    raise ValueError(f'{cores["vermelho"]}Erro opção -> {escolha} <- INVÁLIDA!{cores["semcor"]}')
                                if escolha == 6:
                                    print(f'{cores["negrito"]}VOLTANDO AO MENU INICIAL...{cores["semcor"]}')
                                    sleep(1.8)
                                    break
                            except ValueError as erro:
                                print(erro)
                                continue

                            else:
                                if escolha == 1:
                                    print('\nADICIONAR ALTERAÇÃO\n')
                                    while True:

                                        temp3.clear()
                                        alt = []

                                        with open('Alt.json', 'r') as bd:
                                            dados = json.load(bd)
                                        temp3 = dados.copy()

                                        try:
                                            alteracao = input('Digite o que foi alterado no setor ou sair para voltar\n'
                                                              'ao menu ALTERAÇÕES NO SETOR: ').strip()
                                            if alteracao == 'sair':
                                                print(f'{cores["negrito"]}VOLTANDO AO MENU ALTERAÇÕES NO SETOR...{cores["semcor"]}')
                                                sleep(1.8)
                                                break
                                            if not alteracao:
                                                raise ValueError(f'{cores["vermelho"]}ERRO, você precisa informar algo!{cores["semcor"]}\n')
                                            for alts in temp3:
                                                if alteracao in alts:
                                                    raise ValueError(f'{cores["vermelho"]}ERRO, ESTE EVENTO JÁ FOI REGISTRADO!{cores["semcor"]}\n')
                                        except ValueError as Error:
                                            print(Error)
                                            continue
                                        else:
                                            alt.append(alteracao)
                                            while True:
                                                try:
                                                    data = input('Informe a data do acontecimento no formato dia/mês/ano: ').strip()
                                                    if not data:
                                                        raise ValueError(f'{cores["vermelho"]}ERRO, você precisa informar algo!{cores["semcor"]}\n')
                                                except ValueError as Error:
                                                    print(Error)
                                                    continue
                                                else:
                                                    alt.append(data)
                                                    while True:
                                                        try:
                                                            horario = input('Informe o horário do acontecimento: ').strip()
                                                            if not horario:
                                                                raise ValueError(f'{cores["vermelho"]}ERRO, você precisa informar algo!{cores["semcor"]}\n')
                                                        except ValueError as Error:
                                                            print(Error)
                                                            continue
                                                        else:
                                                            alt.append(horario)
                                                            temp3.append(alt)

                                                            with open('Alt.json','w') as dados:
                                                                json.dump(temp3,dados)
                                                            print(f'{cores["verde"]}Alteração cadastrada COM SUCESSO!{cores["semcor"]}')
                                                            break
                                                    break
                                elif escolha == 2:
                                    print('CONSULTA DE ALTERAÇÕES')
                                    temp3.clear()

                                    with open('Alt.json','r') as Data:
                                        alte = json.load(Data)
                                        temp3 = alte.copy()

                                    if len(temp3) == 0:
                                        print(f'{cores["negrito"]}CADASTRO VAZIO...{cores["semcor"]}')
                                        sleep(1.2)

                                    for cod,alter in enumerate(temp3):
                                        print(f'CÓD:[{cod+1}]- Alteração: {alter[0]:<5} / Data: {alter[1]:^10} / Horário: {alter[2]:>5}')

                                elif escolha == 3:
                                    print('REMOVER ALTERAÇÃO')
                                    while True:
                                        temp3.clear()

                                        with open('Alt.json','r') as altDados:
                                            dados = json.load(altDados)
                                        temp3 = dados.copy()

                                        #Código responsável por deletar registros de alteração
                                        try:
                                            altRem = input('Digite o CÓDIGO da alteração a ser removida ou 0 para voltar ao menu: ').strip()
                                            if not altRem:
                                                raise ValueError(f'{cores["vermelho"]}ERRO, você precisa informar algo!{cores["semcor"]}\n')
                                            elif altRem == '0':
                                                break
                                            elif altRem.isalpha():
                                                raise ValueError(f'{cores["vermelho"]}ERRO, NÃO É PERMITIDO LETRAS!{cores["semcor"]}\n')
                                            else:
                                                altRem = int(altRem)
                                            if altRem > len(temp3) or altRem <=0: #verifica se o número digitado coincide com o indice das listas
                                                raise ValueError(f'{cores["vermelho"]}ERRO, EVENTO NÃO ENCONTRADO!{cores["semcor"]}\n')
                                        except ValueError as Error:
                                            print(Error)
                                            continue
                                        else:
                                            temp3.pop(altRem-1) #precisa ser -1 para coincidir com o indice das listas

                                            with open('Alt.json','w') as altDados:
                                                json.dump(temp3,altDados)
                                            print(f'{cores["verde"]}EVENTO REMOVIDO COM SUCESSO!{cores["semcor"]}')
                                            sleep(1.2)


                                elif escolha == 4:
                                    print(f'{cores["vermelho"]}ATENÇÃO... VOCÊ ESTÁ PRESTES A LIMPAR TODOS OS DADOS CADASTRADOS!{cores["semcor"]}')
                                    while True:
                                        try:
                                            continuar = input('Deseja prosseguir?[s/n]: ').strip()
                                            if not continuar:
                                                raise ValueError(f'{cores["vermelho"]}ERRO, você precisa digitar algo!{cores["semcor"]}\n')
                                        except ValueError as error:
                                            print(error)
                                            continue
                                        else:
                                            if continuar == 'n':
                                                print(f'{cores["negrito"]}Voltando ao MENU...{cores["semcor"]}')
                                                sleep(1.8)
                                                break

                                            temp3.clear()
                                            with open('Alt.json', 'r') as dados:
                                                cadastro = json.load(dados)
                                            temp3 = cadastro.copy()

                                            for alteracao in tqdm(temp3):
                                                sleep(0.5)

                                            with open('Alt.json','w') as AltDados:
                                                emptyData = []
                                                json.dump(emptyData,AltDados)
                                            print(f'{cores["verde"]}Alterações esvaziadas com sucesso!{cores["semcor"]}')
                                            sleep(1.5)
                                            break
                                elif escolha == 5:
                                    print('ENVIAR ALTERAÇÕES POR E-MAIL')

                                    temp3.clear()
                                    with open('Alt.json','r') as bdAlt:
                                        alts = json.load(bdAlt)
                                    temp3 = alts.copy()

                                    if len(temp3) == 0:
                                        print(f'{cores["vermelho"]}CADASTRO DE ALTERAÇÕES VAZIO! RETORNANDO AO MENU ANTERIOR...{cores["semcor"]}')
                                        sleep(1.2)
                                    else:
                                        #Código para criar documento word com uma tabela preenchida com os dados ==============

                                        doc = docx.Document()  # cria um arquivo em formato word
                                        doc.add_heading('Alterações Registradas do Lab. de Informática/Prédio 700', 0)  # Adiciona Título a tabela

                                        # data contém as informações que serão salvas na tabela dentro do arquivo

                                        table = doc.add_table(rows=1, cols=3,style=( 'Colorful List'))  # determina o número e linhas e colunas
                                        titulo_celula = table.rows[0].cells  # armazena as celulas da linha 0 da tabela (corresponde ao título das células)
                                        titulo_celula[0].text = 'Alteração'  # define o título da célula dentro da tabela
                                        titulo_celula[1].text = 'Data Da alteração'  # define o título da célula dentro da tabela
                                        titulo_celula[2].text = 'Horário'


                                        for alt, data, hora in temp3:  # desta linha para baixo, preenche as linhas da tabela com os dados em data.
                                            linha = table.add_row().cells
                                            linha[0].text = alt
                                            linha[1].text = data
                                            linha[2].text = hora


                                        doc.save('Alteracoes.docx')

                                        print('ENVIANDO E-MAIL...')
                                        for dados in tqdm(temp3):
                                            sleep(0.6)

                                        EnviaEmail('Alteracoes.docx')

                                        with open('Alt.json','w') as AltsDados:
                                            emptyData = []
                                            json.dump(emptyData,AltsDados)

                                        sleep(1.6)
                    elif menuin == 5:
                        while True:
                            try:
                                # CÓDIGO PARA CONTROLE DE PROJETORES ======================================
                                escolha = menu(['ADICIONAR PROJETOR','REMOVER PROJETOR','CONSULTAR PROJETORES', 'LIMPAR DADOS','SAIR'],'CONTROLE DE PRJETORES')
                                if escolha < 1 or escolha > 5:
                                    raise ValueError(f'{cores["vermelho"]}ERRO, OPÇÃO -> {escolha} <- inválida!{cores["semcor"]}')
                            except ValueError as erro:
                                print(erro)
                                continue
                            else:
                                if escolha == 5:
                                    print('RETORNANDO AO MENU...')
                                    sleep(1.5)
                                    break
                                elif escolha == 1:
                                    print(f'\n{cores["negrito"]}>>>>>> ADICIONAR PROJETOR <<<<<<<{cores["semcor"]}\n')
                                    while True:
                                        temp4.clear()

                                        with open('Projetores.json','r') as Pdados:
                                            projetores = json.load(Pdados)

                                        temp4 = projetores.copy()
                                        try:

                                            novoProj = input('Digite a nomenclatura do projetor a ser adicionado ou sair para voltar ao menu: ').strip()
                                            if not novoProj:
                                                raise ValueError(f'{cores["vermelho"]}ERRO VOCÊ PRECISA DIGITAR ALGO!{cores["semcor"]}')
                                            elif novoProj =='sair':
                                                print('Voltando ao MENU...')
                                                sleep(1.8)
                                                break
                                            elif novoProj in temp4:
                                                raise ValueError(
                                                    f'{cores["vermelho"]}ERRO, PROJETOR JÁ CADASTRADO!{cores["semcor"]}')
                                        except ValueError as erro:
                                            print(erro)
                                            continue
                                        else:
                                            temp4.append(novoProj)

                                            with open('Projetores.json','w') as Pdados:
                                                json.dump(temp4,Pdados)
                                            print(f'Projetor {novoProj} adicionado com sucesso!!!')
                                elif escolha == 2:
                                    print(
                                        f'\n{cores["negrito"]}>>>>>> REMOVER PROJETOR <<<<<<<{cores["semcor"]}\n')

                                    while True:
                                        temp4.clear()

                                        with open('Projetores.json','r') as Pdados:
                                            projetores = json.load(Pdados)
                                        temp4 = projetores.copy()

                                        try:
                                            removerP = input('Digite qual projetor deseja remover ou sair para voltar ao MENU: ')
                                            if removerP == 'sair':
                                                print('VOLTANDO AO MENU...')
                                                sleep(1.8)
                                                break
                                            elif not removerP:
                                                raise ValueError(f'{cores["vermelho"]}ERRO VOCÊ PRECISA DIGITAR ALGO!{cores["semcor"]}')
                                            elif removerP not in temp4:
                                                raise ValueError(
                                                    f'{cores["vermelho"]}ERRO, PROJETOR NÃO CADASTRADO!{cores["semcor"]}')
                                        except ValueError as erro:
                                            print(erro)
                                            continue
                                        else:
                                            temp4.remove(removerP)

                                            with open('Projetores.json','w') as Pdados:
                                                json.dump(temp4,Pdados)
                                            print(f'{cores["verde"]}Projetor {removerP} removido com SUCESSO!!!{cores["semcor"]}')

                                elif escolha == 3:
                                    print(f'\n{cores["negrito"]}>>>>>> CONSULTA DE PROJETORES DISPONÍVEIS <<<<<<<{cores["semcor"]}\n')
                                    temp4.clear()
                                    with open('Projetores.json','r') as Pdados:
                                        projetores = json.load(Pdados)

                                    temp4 = projetores.copy()

                                    for projetor in temp4:
                                        print(f' - {projetor}')

                                    input('\nAPERTE ENTER PARA VOLTAR AO MENU...')

                                elif escolha == 4:
                                    print(
                                        f'{cores["vermelho"]}ATENÇÃO... VOCÊ ESTÁ PRESTES A LIMPAR TODOS OS DADOS CADASTRADOS!{cores["semcor"]}')
                                    while True:
                                        try:
                                            opc = ['s','n']
                                            continuar = input('Deseja prosseguir?[s/n]: ').strip()
                                            if not continuar:
                                                raise ValueError(
                                                    f'{cores["vermelho"]}ERRO, você precisa digitar algo!{cores["semcor"]}\n')
                                            elif continuar not in opc:
                                                raise ValueError(
                                                    f'{cores["vermelho"]}ERRO, digite s ou n !{cores["semcor"]}\n')
                                        except ValueError as error:
                                            print(error)
                                            continue
                                        else:
                                            if continuar == 'n':
                                                print(f'{cores["negrito"]}Voltando ao MENU...{cores["semcor"]}')
                                                sleep(1.8)
                                                break

                                            temp4.clear()
                                            with open('Projetores.json','r') as pDados:
                                                projetores = json.load(pDados)

                                            print('Esvaziando cadastro...')

                                            for dados in tqdm(projetores):
                                                sleep(0.5)

                                            with open('Projetores.json','w') as pDados:
                                                emptyDados = []
                                                json.dump(emptyDados,pDados)
                                            print('CADASTRO ESVAZIADO COM SUCESSO!')
                                            break
                    elif menuin == 6:
                        print('>>>> COBRANÇA DE MATERIAIS <<<<<')
                        while True:
                            try:
                                material = input('Digite o material a ser cobrado ou sair para voltar ao menu: ').strip()
                                if not material:
                                    raise ValueError(
                                        f'{cores["vermelho"]}ERRO, você precisa digitar algo!{cores["semcor"]}\n')
                                elif material == 'sair':
                                    break
                            except ValueError as error:
                                print(error)
                                continue
                            else:
                                email = input('Digite o e-mail do destinatário: ')
                                EnviaMensgemEmail(email,material)
                                break


        elif opc_inicial == 2:
            micon = AssisVirtual()
            if micon == False:
                pass

        elif opc_inicial == 3:
            pinta('verde', 'OBRIGADO POR UTILIZAR <<<-= O GDP / GESTÃO DE PORTARIA =->>>, FINALIZANDO... ')
            sleep(3.0)
            break
